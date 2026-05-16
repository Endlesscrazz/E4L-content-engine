"""
Doc-type-aware chunkers for the E4L source corpus.

Chunking strategy is a citation-integrity decision, not a performance one:
chunk boundaries map to natural semantic units so every retrievable unit
equals one citable claim. Four doc types, four strategies:

  narrative    — Origin Story, Differentiation: blank-separated paragraph groups
  ai_version   — AI Version book: section-title detection (blank → short line → new chunk)
  product_doc  — miHealth, BWS: ___ / *** separator lines
  research_finding — Research Summary txt: ___ separator lines, one finding per chunk

Each chunker returns List[SourceChunk] with chunk_id + content populated.
Embeddings and corpus annotations are added by scripts/ingest_corpus.py.
"""
from __future__ import annotations

import re
from pathlib import Path

from docx import Document

from app.models import SourceChunk


def _slug(doc_name: str) -> str:
    name = Path(doc_name).stem.lower()
    name = re.sub(r"[^a-z0-9]+", "_", name).strip("_")
    return name[:30]


def _chunk_id(slug: str, index: int) -> str:
    return f"{slug}_{index:04d}"


# ── Narrative chunker (Origin Story, Differentiation) ─────────────────────────

def chunk_narrative(doc_path: Path) -> list[SourceChunk]:
    """
    Groups consecutive non-empty paragraphs separated by blank lines.
    One blank-separated block = one citable chunk.
    Short header-like lines are kept in the same block as the following content
    since they are not separated by a blank — they belong together.
    """
    doc = Document(doc_path)
    doc_name = doc_path.name
    slug = _slug(doc_name)

    blocks: list[list[str]] = []
    current: list[str] = []

    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            current.append(text)
        else:
            if current:
                blocks.append(current)
                current = []
    if current:
        blocks.append(current)

    return [
        SourceChunk(
            chunk_id=_chunk_id(slug, i),
            doc_name=doc_name,
            doc_type="narrative",
            content="\n".join(block),
        )
        for i, block in enumerate(blocks)
    ]


# ── AI Version chunker (section-based) ────────────────────────────────────────

def chunk_ai_version(doc_path: Path) -> list[SourceChunk]:
    """
    Section-based chunking for the AI Version book.

    A paragraph is treated as a section title when:
      - it follows at least one blank / whitespace-only line (or is the first para)
      - its stripped length is <= 70 chars
      - it does not start with a digit (excludes numbered list items)
      - it does not start with * or - (excludes bullet points)

    All paragraphs following a title accumulate into that section's chunk
    until the next detected title.
    """
    doc = Document(doc_path)
    doc_name = doc_path.name
    slug = _slug(doc_name)

    paragraphs = [p.text for p in doc.paragraphs]
    sections: list[list[str]] = []
    current: list[str] = []
    prev_was_blank = True  # treat doc start as if preceded by blank

    for text in paragraphs:
        stripped = text.strip()
        if not stripped:
            prev_was_blank = True
            continue

        is_title = (
            prev_was_blank
            and len(stripped) <= 70
            and not stripped[0].isdigit()
            and not stripped.startswith("*")
            and not stripped.startswith("-")
        )

        if is_title and current:
            sections.append(current)
            current = [stripped]
        else:
            current.append(stripped)

        prev_was_blank = False

    if current:
        sections.append(current)

    return [
        SourceChunk(
            chunk_id=_chunk_id(slug, i),
            doc_name=doc_name,
            doc_type="ai_version",
            content="\n".join(section),
        )
        for i, section in enumerate(sections)
        if any(s.strip() for s in section)
    ]


# ── Product doc chunker (miHealth, BWS) ───────────────────────────────────────

_SEPARATOR_RE = re.compile(r"^[_*]{3,}$")


def chunk_product_doc(doc_path: Path) -> list[SourceChunk]:
    """
    Splits on separator lines (3+ underscores or asterisks).
    miHealth uses _______, BWS uses ***.
    Everything between separators = one chunk.
    """
    doc = Document(doc_path)
    doc_name = doc_path.name
    slug = _slug(doc_name)

    sections: list[list[str]] = []
    current: list[str] = []

    for para in doc.paragraphs:
        text = para.text.strip()
        if _SEPARATOR_RE.match(text):
            if current:
                sections.append(current)
                current = []
        elif text:
            current.append(text)

    if current:
        sections.append(current)

    return [
        SourceChunk(
            chunk_id=_chunk_id(slug, i),
            doc_name=doc_name,
            doc_type="product_doc",
            content="\n".join(section),
        )
        for i, section in enumerate(sections)
        if any(s.strip() for s in section)
    ]


# ── Research Summary chunker (per-finding) ────────────────────────────────────

def chunk_research_summary(txt_path: Path) -> list[SourceChunk]:
    """
    Splits on separator lines (8+ underscores).
    Each finding block = study title + supporting paragraphs = one chunk.
    One chunk = one citable research finding, supporting exact citation.
    """
    doc_name = txt_path.name
    slug = _slug(doc_name)
    raw = txt_path.read_text(encoding="utf-8")

    raw_sections = re.split(r"\n_{8,}\n", raw)

    chunks = []
    for i, section in enumerate(raw_sections):
        content = section.strip()
        if content:
            chunks.append(
                SourceChunk(
                    chunk_id=_chunk_id(slug, i),
                    doc_name=doc_name,
                    doc_type="research_finding",
                    content=content,
                )
            )
    return chunks
